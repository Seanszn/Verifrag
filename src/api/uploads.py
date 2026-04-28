"""Upload routes for server-ingested user documents."""

from __future__ import annotations

import io
import json
import re
from pathlib import Path
from typing import Any

from fastapi import APIRouter, Depends, File, Form, HTTPException, UploadFile, status
from pydantic import BaseModel

from src.api.dependencies import get_current_user
from src.config import DATA_DIR
from src.ingestion.chunker import chunk_document
from src.ingestion.document import LegalDocument
from src.ingestion.pdf_parser import PDFParserError, parse_pdf_to_document
from src.ingestion.user_documents import (
    UserDocumentError,
    decode_text_bytes,
    make_user_document_id,
    upsert_jsonl_rows,
    user_document_to_row,
)
from src.retrieval.user_uploads import build_user_upload_indices
from src.indexing.index_discovery import discover_index_artifacts

try:
    from docx import Document as DocxDocument
except ImportError:  # pragma: no cover - depends on optional runtime package
    DocxDocument = None


router = APIRouter(tags=["uploads"])

USER_UPLOADS_ROOT = DATA_DIR / "uploads"
ALLOWED_UPLOAD_SUFFIXES = {".pdf", ".txt", ".md", ".docx"}
MAX_UPLOADS_PER_REQUEST = 10
MAX_UPLOAD_BYTES = 25 * 1024 * 1024
RAW_FILENAME = "user_uploads.jsonl"
PROCESSED_FILENAME = "user_upload_chunks.jsonl"


class UploadFileSummary(BaseModel):
    """Metadata for one uploaded file."""

    filename: str
    content_type: str | None
    document_id: str
    size_bytes: int
    chunk_count: int
    is_privileged: bool


class UploadResponse(BaseModel):
    """Server response for upload ingestion."""

    conversation_id: int | None
    files_uploaded: int
    documents_upserted: int
    chunks_upserted: int
    files: list[UploadFileSummary]


class UnsupportedUploadError(ValueError):
    """Raised when an uploaded file cannot be processed."""


@router.post(
    "/api/uploads",
    response_model=UploadResponse,
    status_code=status.HTTP_201_CREATED,
)
async def upload_documents(
    files: list[UploadFile] = File(...),
    conversation_id: int | None = Form(default=None),
    is_privileged: bool = Form(default=True),
    current_user: dict[str, Any] = Depends(get_current_user),
) -> UploadResponse:
    """Receive user files, parse them into documents, and persist chunk JSONL."""
    if not files:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="No files were uploaded.")
    if len(files) > MAX_UPLOADS_PER_REQUEST:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Upload at most {MAX_UPLOADS_PER_REQUEST} files per request.",
        )

    parsed_uploads: list[dict[str, Any]] = []
    for uploaded_file in files:
        filename = _sanitize_filename(uploaded_file.filename)
        try:
            file_bytes = await uploaded_file.read()
            _validate_upload(filename, file_bytes)
            document = _document_from_upload(file_bytes, filename, is_privileged=is_privileged)
            chunks = chunk_document(document)
        except UnsupportedUploadError as exc:
            raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail=str(exc)) from exc
        except PDFParserError as exc:
            raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail=str(exc)) from exc
        finally:
            await uploaded_file.close()

        parsed_uploads.append(
            {
                "filename": filename,
                "content_type": uploaded_file.content_type,
                "size_bytes": len(file_bytes),
                "document": document,
                "chunks": chunks,
                "file_bytes": file_bytes,
            }
        )

    user_root = USER_UPLOADS_ROOT / f"user_{current_user['id']}"
    raw_path = user_root / "raw" / RAW_FILENAME
    processed_path = user_root / "processed" / PROCESSED_FILENAME
    stored_files_dir = user_root / "files"

    raw_rows = [user_document_to_row(item["document"]) for item in parsed_uploads]
    processed_rows = [chunk.to_dict() for item in parsed_uploads for chunk in item["chunks"]]

    for item in parsed_uploads:
        _store_original_file(
            stored_files_dir=stored_files_dir,
            filename=item["filename"],
            document_id=item["document"].id,
            file_bytes=item["file_bytes"],
        )

    upsert_jsonl_rows(raw_path, raw_rows)
    upsert_jsonl_rows(processed_path, processed_rows)
    build_user_upload_indices(
        current_user["id"],
        uploads_root=USER_UPLOADS_ROOT,
    )

    return UploadResponse(
        conversation_id=conversation_id,
        files_uploaded=len(parsed_uploads),
        documents_upserted=len(raw_rows),
        chunks_upserted=len(processed_rows),
        files=[
            UploadFileSummary(
                filename=item["filename"],
                content_type=item["content_type"],
                document_id=item["document"].id,
                size_bytes=item["size_bytes"],
                chunk_count=len(item["chunks"]),
                is_privileged=item["document"].is_privileged,
            )
            for item in parsed_uploads
        ],
    )


class UploadedDocumentInfo(BaseModel):
    """Information about a user-uploaded document."""

    document_id: str
    filename: str
    case_name: str
    size_bytes: int
    chunk_count: int
    is_privileged: bool
    uploaded_at: str


@router.get(
    "/api/uploads",
    response_model=list[UploadedDocumentInfo],
)
async def list_uploads(
    current_user: dict[str, Any] = Depends(get_current_user),
) -> list[UploadedDocumentInfo]:
    """List all uploaded documents for the current user."""
    user_root = USER_UPLOADS_ROOT / f"user_{current_user['id']}"
    raw_path = user_root / "raw" / RAW_FILENAME
    processed_path = user_root / "processed" / PROCESSED_FILENAME

    if not raw_path.exists():
        return []

    document_ids: set[str] = set()
    with raw_path.open("r", encoding="utf-8") as f:
        for line in f:
            if line.strip():
                try:
                    doc = json.loads(line)
                    document_ids.add(doc["id"])
                except (json.JSONDecodeError, KeyError):
                    continue

    processed_counts: dict[str, int] = {}
    if processed_path.exists():
        with processed_path.open("r", encoding="utf-8") as f:
            for line in f:
                if line.strip():
                    try:
                        chunk = json.loads(line)
                        doc_id = chunk.get("doc_id") or chunk.get("document_id")
                        if doc_id:
                            processed_counts[doc_id] = processed_counts.get(doc_id, 0) + 1
                    except (json.JSONDecodeError, KeyError):
                        continue

    documents: list[UploadedDocumentInfo] = []
    with raw_path.open("r", encoding="utf-8") as f:
        for line in f:
            if line.strip():
                try:
                    doc = json.loads(line)
                    doc_id = doc.get("id")
                    if doc_id in document_ids:
                        document_ids.discard(doc_id)
                        documents.append(
                            UploadedDocumentInfo(
                                document_id=doc_id,
                                filename=doc.get("source_file", ""),
                                case_name=doc.get("case_name", ""),
                                size_bytes=len(doc.get("full_text", "")),
                                chunk_count=processed_counts.get(doc_id, 0),
                                is_privileged=doc.get("is_privileged", False),
                                uploaded_at=doc.get("uploaded_at", ""),
                            )
                        )
                except (json.JSONDecodeError, KeyError):
                    continue

    return documents


@router.delete(
    "/api/uploads/{document_id}",
    status_code=status.HTTP_204_NO_CONTENT,
)
async def delete_upload(
    document_id: str,
    current_user: dict[str, Any] = Depends(get_current_user),
) -> None:
    """Delete an uploaded document for the current user."""
    user_root = USER_UPLOADS_ROOT / f"user_{current_user['id']}"
    raw_path = user_root / "raw" / RAW_FILENAME
    processed_path = user_root / "processed" / PROCESSED_FILENAME
    index_dir = user_root / "index"

    if not raw_path.exists():
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="No uploads found.",
        )

    remaining_raw: list[dict[str, Any]] = []
    deleted = False
    with raw_path.open("r", encoding="utf-8") as f:
        for line in f:
            if line.strip():
                try:
                    doc = json.loads(line)
                    if doc.get("id") == document_id:
                        deleted = True
                        continue
                    remaining_raw.append(doc)
                except json.JSONDecodeError:
                    continue

    if not deleted:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Document not found.",
        )

    with raw_path.open("w", encoding="utf-8") as f:
        for doc in remaining_raw:
            f.write(json.dumps(doc, ensure_ascii=False) + "\n")

    remaining_processed: list[dict[str, Any]] = []
    if processed_path.exists():
        with processed_path.open("r", encoding="utf-8") as f:
            for line in f:
                if line.strip():
                    try:
                        chunk = json.loads(line)
                        chunk_doc_id = chunk.get("doc_id") or chunk.get("document_id")
                        if chunk_doc_id != document_id:
                            remaining_processed.append(chunk)
                    except json.JSONDecodeError:
                        continue

    with processed_path.open("w", encoding="utf-8") as f:
        for chunk in remaining_processed:
            f.write(json.dumps(chunk, ensure_ascii=False) + "\n")

    if index_dir.exists():
        import shutil
        try:
            artifacts = discover_index_artifacts(index_dir=index_dir)
            if artifacts.chroma_path.exists():
                shutil.rmtree(artifacts.chroma_path)
            if artifacts.bm25_path.exists():
                shutil.rmtree(artifacts.bm25_path)
        except Exception:
            pass

    build_user_upload_indices(
        current_user["id"],
        uploads_root=USER_UPLOADS_ROOT,
    )


def _validate_upload(filename: str, file_bytes: bytes) -> None:
    if not filename:
        raise UnsupportedUploadError("Uploaded files must have a filename.")
    suffix = Path(filename).suffix.lower()
    if suffix not in ALLOWED_UPLOAD_SUFFIXES:
        allowed = ", ".join(sorted(ALLOWED_UPLOAD_SUFFIXES))
        raise UnsupportedUploadError(f"Unsupported file type '{suffix or '<none>'}'. Allowed: {allowed}.")
    if not file_bytes:
        raise UnsupportedUploadError(f"Uploaded file '{filename}' is empty.")
    if len(file_bytes) > MAX_UPLOAD_BYTES:
        raise UnsupportedUploadError(
            f"Uploaded file '{filename}' exceeds the {MAX_UPLOAD_BYTES} byte limit."
        )


def _document_from_upload(file_bytes: bytes, filename: str, *, is_privileged: bool) -> LegalDocument:
    suffix = Path(filename).suffix.lower()
    if suffix == ".pdf":
        return parse_pdf_to_document(
            file_bytes,
            filename=filename,
            is_privileged=is_privileged,
        )
    if suffix in {".txt", ".md"}:
        return _text_document_from_bytes(file_bytes, filename, is_privileged=is_privileged)
    if suffix == ".docx":
        return _docx_document_from_bytes(file_bytes, filename, is_privileged=is_privileged)
    raise UnsupportedUploadError(f"Unsupported file type '{suffix}'.")


def _text_document_from_bytes(
    file_bytes: bytes,
    filename: str,
    *,
    is_privileged: bool,
) -> LegalDocument:
    try:
        text = decode_text_bytes(
            file_bytes,
            error_message="Text upload could not be decoded as UTF-8 or a common fallback encoding.",
        ).strip()
    except UserDocumentError as exc:
        raise UnsupportedUploadError(str(exc)) from exc
    if not text:
        raise UnsupportedUploadError(f"Uploaded text file '{filename}' did not contain readable text.")
    return LegalDocument(
        id=make_user_document_id(filename, text),
        doc_type="user_upload",
        full_text=text,
        case_name=Path(filename).stem.replace("_", " ").replace("-", " ").strip() or filename,
        source_file=filename,
        is_privileged=is_privileged,
    )


def _docx_document_from_bytes(
    file_bytes: bytes,
    filename: str,
    *,
    is_privileged: bool,
) -> LegalDocument:
    if DocxDocument is None:
        raise UnsupportedUploadError("DOCX uploads require the python-docx package.")

    try:
        doc = DocxDocument(io.BytesIO(file_bytes))
    except Exception as exc:  # pragma: no cover - parser-specific failure mode
        raise UnsupportedUploadError(f"Failed to parse DOCX file '{filename}'.") from exc

    paragraphs = [paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text.strip()]
    text = "\n\n".join(paragraphs).strip()
    if not text:
        raise UnsupportedUploadError(f"Uploaded DOCX file '{filename}' did not contain readable text.")

    return LegalDocument(
        id=make_user_document_id(filename, text),
        doc_type="user_upload",
        full_text=text,
        case_name=Path(filename).stem.replace("_", " ").replace("-", " ").strip() or filename,
        source_file=filename,
        is_privileged=is_privileged,
    )


def _store_original_file(
    *,
    stored_files_dir: Path,
    filename: str,
    document_id: str,
    file_bytes: bytes,
) -> Path:
    stored_files_dir.mkdir(parents=True, exist_ok=True)
    suffix = Path(filename).suffix.lower()
    stored_path = stored_files_dir / f"{document_id}{suffix}"
    stored_path.write_bytes(file_bytes)
    return stored_path


def _sanitize_filename(filename: str | None) -> str:
    if not filename:
        return ""
    basename = Path(filename).name.strip()
    sanitized = re.sub(r"[^A-Za-z0-9._ -]+", "_", basename)
    return sanitized.strip(" .")


