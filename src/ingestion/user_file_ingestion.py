"""
Helpers for ingesting user-supplied files into the local corpus.
"""

from __future__ import annotations

import hashlib
import json
import logging
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Iterable, Sequence

from src.config import PROCESSED_DIR, RAW_DIR
from src.ingestion.chunker import chunk_document
from src.ingestion.document import LegalDocument
from src.ingestion.pdf_parser import PDFParser

try:
    from docx import Document as DocxDocument
except ImportError:  # pragma: no cover - depends on environment
    DocxDocument = None

logger = logging.getLogger(__name__)

_TEXT_SUFFIXES = {".txt", ".md", ".markdown", ".rst"}
_DOCX_SUFFIXES = {".docx"}
_PDF_SUFFIXES = {".pdf"}
_SUPPORTED_SUFFIXES = _TEXT_SUFFIXES | _DOCX_SUFFIXES | _PDF_SUFFIXES
_MULTISPACE_RE = re.compile(r"[ \t]+")


class UnsupportedUserFileError(ValueError):
    """Raised when a user-supplied file type is not supported."""


@dataclass
class UserFileIngestionSummary:
    """High-level summary of one user-file ingestion run."""

    files_discovered: int = 0
    files_ingested: int = 0
    documents_upserted: int = 0
    chunks_upserted: int = 0
    raw_output_path: Path | None = None
    processed_output_path: Path | None = None
    document_ids: list[str] = field(default_factory=list)


class UserFileCorpusIngestor:
    """Parse local files and persist user-upload documents/chunks as JSONL."""

    def __init__(
        self,
        *,
        raw_dir: Path | str = RAW_DIR,
        processed_dir: Path | str = PROCESSED_DIR,
        raw_output_file: str = "user_uploads.jsonl",
        processed_output_file: str = "user_upload_chunks.jsonl",
        pdf_parser: PDFParser | None = None,
    ) -> None:
        self.raw_dir = Path(raw_dir)
        self.processed_dir = Path(processed_dir)
        self.raw_output_path = self.raw_dir / raw_output_file
        self.processed_output_path = self.processed_dir / processed_output_file
        self.pdf_parser = pdf_parser or PDFParser()

        self.raw_dir.mkdir(parents=True, exist_ok=True)
        self.processed_dir.mkdir(parents=True, exist_ok=True)

    @property
    def supported_suffixes(self) -> frozenset[str]:
        return frozenset(_SUPPORTED_SUFFIXES)

    def ingest_path(
        self,
        path: str | Path,
        *,
        is_privileged: bool = False,
    ) -> UserFileIngestionSummary:
        """Convenience wrapper for ingesting a single file."""
        return self.ingest_paths([path], is_privileged=is_privileged)

    def ingest_paths(
        self,
        paths: Iterable[str | Path],
        *,
        recursive: bool = False,
        is_privileged: bool = False,
    ) -> UserFileIngestionSummary:
        """Parse and persist one or more user files into raw/processed JSONL outputs."""
        resolved_files = self._resolve_input_files(paths, recursive=recursive)
        documents_by_id = {
            record["id"]: record
            for record in self._load_jsonl_records(self.raw_output_path)
            if record.get("id")
        }
        chunks_by_id = {
            record["id"]: record
            for record in self._load_jsonl_records(self.processed_output_path)
            if record.get("id")
        }

        summary = UserFileIngestionSummary(
            files_discovered=len(resolved_files),
            raw_output_path=self.raw_output_path,
            processed_output_path=self.processed_output_path,
        )

        for path in resolved_files:
            document = self.parse_file(path, is_privileged=is_privileged)
            chunks = chunk_document(document)

            documents_by_id[document.id] = _document_to_record(document, source_path=path)
            for chunk_id in [
                existing_id
                for existing_id, record in chunks_by_id.items()
                if record.get("doc_id") == document.id
            ]:
                del chunks_by_id[chunk_id]

            for chunk in chunks:
                chunks_by_id[chunk.id] = chunk.to_dict()

            summary.files_ingested += 1
            summary.documents_upserted += 1
            summary.chunks_upserted += len(chunks)
            summary.document_ids.append(document.id)

        self._write_jsonl_records(self.raw_output_path, documents_by_id.values())
        self._write_jsonl_records(self.processed_output_path, chunks_by_id.values())
        return summary

    def parse_file(
        self,
        path: str | Path,
        *,
        is_privileged: bool = False,
    ) -> LegalDocument:
        """Parse a supported file on disk into a LegalDocument."""
        resolved = Path(path)
        suffix = resolved.suffix.lower()

        if suffix in _PDF_SUFFIXES:
            return self.pdf_parser.parse(resolved, is_privileged=is_privileged)

        if suffix in _DOCX_SUFFIXES:
            return self._parse_docx(resolved, is_privileged=is_privileged)

        if suffix in _TEXT_SUFFIXES:
            return self._parse_text(resolved, is_privileged=is_privileged)

        raise UnsupportedUserFileError(
            f"Unsupported file type: {resolved.suffix or '<no extension>'}. "
            f"Supported types: {', '.join(sorted(self.supported_suffixes))}"
        )

    def _resolve_input_files(
        self,
        paths: Iterable[str | Path],
        *,
        recursive: bool,
    ) -> list[Path]:
        resolved_files: list[Path] = []

        for raw_path in paths:
            path = Path(raw_path)
            if not path.exists():
                raise FileNotFoundError(f"Input path does not exist: {path}")

            if path.is_file():
                self._validate_supported_file(path)
                resolved_files.append(path)
                continue

            pattern = "**/*" if recursive else "*"
            for candidate in sorted(item for item in path.glob(pattern) if item.is_file()):
                if candidate.suffix.lower() not in _SUPPORTED_SUFFIXES:
                    continue
                resolved_files.append(candidate)

        if not resolved_files:
            raise FileNotFoundError("No supported files were found in the provided paths.")

        unique_files = []
        seen: set[Path] = set()
        for item in resolved_files:
            resolved_item = item.resolve()
            if resolved_item in seen:
                continue
            seen.add(resolved_item)
            unique_files.append(item)
        return unique_files

    def _validate_supported_file(self, path: Path) -> None:
        if path.suffix.lower() not in _SUPPORTED_SUFFIXES:
            raise UnsupportedUserFileError(
                f"Unsupported file type: {path.suffix or '<no extension>'}. "
                f"Supported types: {', '.join(sorted(self.supported_suffixes))}"
            )

    def _parse_text(self, path: Path, *, is_privileged: bool) -> LegalDocument:
        text = path.read_text(encoding="utf-8")
        normalized = _normalize_text(text)
        if not normalized:
            raise ValueError(f"File does not contain any readable text: {path}")

        return LegalDocument(
            id=_make_document_id(path.name, normalized),
            doc_type="user_upload",
            full_text=normalized,
            case_name=_infer_title_from_text(path.stem, normalized),
            source_file=path.name,
            is_privileged=is_privileged,
        )

    def _parse_docx(self, path: Path, *, is_privileged: bool) -> LegalDocument:
        if DocxDocument is None:
            raise RuntimeError(
                "python-docx is not installed. Install `python-docx` to enable DOCX ingestion."
            )

        document = DocxDocument(str(path))
        paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
        text = "\n\n".join(paragraphs)
        normalized = _normalize_text(text)
        if not normalized:
            raise ValueError(f"File does not contain any readable text: {path}")

        return LegalDocument(
            id=_make_document_id(path.name, normalized),
            doc_type="user_upload",
            full_text=normalized,
            case_name=_infer_title_from_text(path.stem, normalized),
            source_file=path.name,
            is_privileged=is_privileged,
        )

    @staticmethod
    def _load_jsonl_records(path: Path) -> list[dict]:
        if not path.exists():
            return []

        records: list[dict] = []
        with path.open("r", encoding="utf-8") as handle:
            for line in handle:
                line = line.strip()
                if not line:
                    continue
                try:
                    records.append(json.loads(line))
                except json.JSONDecodeError:
                    logger.warning("Skipping invalid JSONL line in %s", path)
        return records

    @staticmethod
    def _write_jsonl_records(path: Path, records: Sequence[dict]) -> None:
        sorted_records = sorted(
            (dict(record) for record in records),
            key=lambda item: (
                str(item.get("source_file") or ""),
                str(item.get("doc_id") or item.get("id") or ""),
                str(item.get("id") or ""),
            ),
        )
        with path.open("w", encoding="utf-8") as handle:
            for record in sorted_records:
                handle.write(json.dumps(record, ensure_ascii=False) + "\n")


def _document_to_record(document: LegalDocument, *, source_path: Path) -> dict:
    return {
        "id": document.id,
        "doc_type": document.doc_type,
        "case_name": document.case_name,
        "citation": document.citation,
        "court": document.court,
        "court_level": document.court_level,
        "date_decided": str(document.date_decided) if document.date_decided else None,
        "full_text": document.full_text,
        "source_file": document.source_file,
        "is_privileged": document.is_privileged,
        "source_path": str(source_path.resolve()),
    }


def _normalize_text(text: str) -> str:
    normalized = text.replace("\r\n", "\n").replace("\r", "\n")
    normalized = _MULTISPACE_RE.sub(" ", normalized)
    normalized = re.sub(r"[ \t]*\n[ \t]*", "\n", normalized)
    normalized = re.sub(r"\n{3,}", "\n\n", normalized)
    return normalized.strip()


def _infer_title_from_text(fallback_name: str, text: str) -> str:
    for line in text.splitlines():
        candidate = line.strip()
        if len(candidate) < 6:
            continue
        if len(candidate) > 160:
            continue
        return candidate
    return fallback_name.replace("_", " ").replace("-", " ").strip() or fallback_name


def _make_document_id(source_name: str, text: str) -> str:
    slug = re.sub(r"[^a-z0-9]+", "-", Path(source_name).stem.lower()).strip("-") or "upload"
    digest = hashlib.sha1(text[:4000].encode("utf-8")).hexdigest()[:10]
    return f"upload_{slug}_{digest}"


__all__ = [
    "UnsupportedUserFileError",
    "UserFileCorpusIngestor",
    "UserFileIngestionSummary",
]
